"""Document reader MCP server — extract text and structure from files.

Supports PDF, DOCX, Excel, CSV, draw.io XML, SVG, and plain text.
Document parsing libraries are optional — tools that need missing
libraries return a clear install instruction instead of failing silently.

See ``design/details/document-reader-mcp.md``.
"""

from __future__ import annotations

import csv
import os
import xml.etree.ElementTree as ET
from html import unescape
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP

server = FastMCP("swarmkit-docs-reader")

_workspace_root: Path | None = None

_MAX_TEXT_BYTES = 500 * 1024
_MAX_PDF_PAGES = 50
_MAX_EXCEL_ROWS = 100
_MAX_CSV_ROWS = 100


def _resolve_path(path: str) -> Path:
    """Resolve a path relative to workspace root or as absolute."""
    p = Path(path)
    if p.is_absolute():
        return p
    root = _workspace_root or Path.cwd()
    return root / p


def _truncation_notice(actual: int, limit: int, unit: str, hint: str) -> str:
    return (
        f"\n\n--- TRUNCATED: showing {limit} of {actual} {unit}. "
        f"Use {hint} to read a specific range. ---"
    )


@server.tool()
def read_pdf(
    path: str,
    start_page: int = 1,
    end_page: int | None = None,
) -> str:
    """Extract text from a PDF file.

    Args:
        path: File path (relative to workspace root or absolute).
        start_page: First page to extract (1-indexed, default 1).
        end_page: Last page to extract (inclusive). Defaults to start_page + 49.
    """
    try:
        import pymupdf  # type: ignore[import-not-found]  # noqa: PLC0415
    except ImportError:
        return (
            "ERROR: pymupdf is not installed. Install it with:\n"
            "  pip install pymupdf\n"
            "  # or: uv pip install pymupdf"
        )

    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        doc = pymupdf.open(str(resolved))
    except Exception as exc:
        return f"ERROR: cannot open PDF: {exc}"

    total_pages = len(doc)
    start = max(0, start_page - 1)
    end = min(total_pages, (end_page or start + _MAX_PDF_PAGES))
    truncated = end < total_pages and end_page is None

    parts: list[str] = [f"# {resolved.name} ({total_pages} pages)\n"]
    for i in range(start, end):
        page = doc[i]
        text = page.get_text()
        parts.append(f"## Page {i + 1}\n\n{text.strip()}\n")

    doc.close()

    result = "\n".join(parts)
    if truncated:
        result += _truncation_notice(total_pages, end - start, "pages", "start_page/end_page")
    return result


@server.tool()
def read_docx(path: str) -> str:
    """Extract text from a DOCX file, preserving heading hierarchy.

    Args:
        path: File path (relative to workspace root or absolute).
    """
    try:
        import docx  # type: ignore[import-not-found]  # noqa: PLC0415
    except ImportError:
        return (
            "ERROR: python-docx is not installed. Install it with:\n"
            "  pip install python-docx\n"
            "  # or: uv pip install python-docx"
        )

    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        document = docx.Document(str(resolved))
    except Exception as exc:
        return f"ERROR: cannot open DOCX: {exc}"

    parts: list[str] = [f"# {resolved.name}\n"]
    for para in document.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("Heading 1"):
            parts.append(f"\n## {text}\n")
        elif style.startswith("Heading 2"):
            parts.append(f"\n### {text}\n")
        elif style.startswith("Heading 3"):
            parts.append(f"\n#### {text}\n")
        elif style.startswith("Heading"):
            parts.append(f"\n##### {text}\n")
        else:
            parts.append(text)

    for table in document.tables:
        parts.append(_docx_table_to_markdown(table))

    result = "\n".join(parts)
    if len(result.encode()) > _MAX_TEXT_BYTES:
        result = result[:_MAX_TEXT_BYTES] + "\n\n--- TRUNCATED at 500KB ---"
    return result


def _docx_table_to_markdown(table: Any) -> str:
    """Convert a python-docx Table to a markdown table."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return f"\n{header}\n{separator}\n{body}\n"


@server.tool()
def read_excel(
    path: str,
    sheet: str | None = None,
    max_rows: int = _MAX_EXCEL_ROWS,
) -> str:
    """Extract data from an Excel file as markdown tables.

    Args:
        path: File path (relative to workspace root or absolute).
        sheet: Sheet name to read. Reads all sheets if not specified.
        max_rows: Maximum rows per sheet (default 100).
    """
    try:
        import openpyxl  # type: ignore[import-untyped]  # noqa: PLC0415
    except ImportError:
        return (
            "ERROR: openpyxl is not installed. Install it with:\n"
            "  pip install openpyxl\n"
            "  # or: uv pip install openpyxl"
        )

    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        wb = openpyxl.load_workbook(str(resolved), read_only=True, data_only=True)
    except Exception as exc:
        return f"ERROR: cannot open Excel file: {exc}"

    sheet_names = [sheet] if sheet else wb.sheetnames
    parts: list[str] = [f"# {resolved.name}\n"]

    for sname in sheet_names:
        if sname not in wb.sheetnames:
            parts.append(f"\n## Sheet: {sname}\n\nERROR: sheet not found.\n")
            continue

        ws = wb[sname]
        parts.append(f"\n## Sheet: {sname}\n")

        rows: list[list[str]] = []
        total_rows = 0
        for row in ws.iter_rows(values_only=True):
            total_rows += 1
            if total_rows <= max_rows + 1:
                rows.append([str(c) if c is not None else "" for c in row])

        if not rows:
            parts.append("(empty sheet)\n")
            continue

        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
        body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
        parts.append(f"{header}\n{separator}\n{body}\n")

        if total_rows > max_rows + 1:
            parts.append(_truncation_notice(total_rows - 1, max_rows, "rows", "max_rows"))

    wb.close()
    return "\n".join(parts)


@server.tool()
def read_csv(
    path: str,
    max_rows: int = _MAX_CSV_ROWS,
    delimiter: str = ",",
) -> str:
    """Read a CSV file and return it as a markdown table.

    Args:
        path: File path (relative to workspace root or absolute).
        max_rows: Maximum data rows to include (default 100).
        delimiter: Column delimiter (default comma).
    """
    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        with open(resolved, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows: list[list[str]] = []
            total = 0
            for row in reader:
                total += 1
                if total <= max_rows + 1:
                    rows.append([c.strip() for c in row])
    except Exception as exc:
        return f"ERROR: cannot read CSV: {exc}"

    if not rows:
        return f"# {resolved.name}\n\n(empty file)"

    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    result = f"# {resolved.name}\n\n{header}\n{separator}\n{body}"

    if total > max_rows + 1:
        result += _truncation_notice(total - 1, max_rows, "rows", "max_rows")
    return result


@server.tool()
def read_drawio(path: str) -> str:
    """Parse a draw.io XML file and describe its nodes and connections.

    Works with both .drawio and .drawio.xml files. Extracts shapes,
    labels, and connections into a structured text description.

    Args:
        path: File path (relative to workspace root or absolute).
    """
    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        tree = ET.parse(resolved)
    except ET.ParseError as exc:
        return f"ERROR: cannot parse XML: {exc}"

    root = tree.getroot()
    return _parse_drawio_xml(root, resolved.name)


def _parse_drawio_xml(root: ET.Element, filename: str) -> str:
    """Extract nodes and edges from draw.io XML."""
    parts: list[str] = [f"# {filename} (draw.io diagram)\n"]

    cells = root.findall(".//mxCell") or root.findall(".//{*}mxCell")
    if not cells:
        return f"# {filename}\n\n(no diagram elements found)"

    nodes: dict[str, str] = {}
    edges: list[tuple[str, str, str]] = []

    for cell in cells:
        cell_id = cell.get("id", "")
        value = cell.get("value", "")
        label = unescape(value).strip() if value else ""
        source = cell.get("source", "")
        target = cell.get("target", "")
        is_edge = cell.get("edge") == "1"

        if is_edge and source and target:
            edges.append((source, target, label))
        elif label and cell_id:
            nodes[cell_id] = label

    if nodes:
        parts.append("## Nodes\n")
        for node_id, label in nodes.items():
            parts.append(f"- **{label}** (id: {node_id})")

    if edges:
        parts.append("\n## Connections\n")
        for source, target, label in edges:
            src_label = nodes.get(source, source)
            tgt_label = nodes.get(target, target)
            arrow = f" —[{label}]→ " if label else " → "
            parts.append(f"- {src_label}{arrow}{tgt_label}")

    return "\n".join(parts)


@server.tool()
def read_svg(path: str) -> str:
    """Parse an SVG file and extract text elements and structure.

    Useful for architecture diagrams and flowcharts saved as SVG.

    Args:
        path: File path (relative to workspace root or absolute).
    """
    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        tree = ET.parse(resolved)
    except ET.ParseError as exc:
        return f"ERROR: cannot parse SVG: {exc}"

    root = tree.getroot()
    ns = {"svg": "http://www.w3.org/2000/svg"}

    texts: list[str] = []
    for elem in root.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "text":
            t = "".join(elem.itertext()).strip()
            if t:
                texts.append(t)
        elif tag == "title":
            t = (elem.text or "").strip()
            if t:
                texts.append(f"[title: {t}]")
        elif tag == "desc":
            t = (elem.text or "").strip()
            if t:
                texts.append(f"[description: {t}]")

    title_elem = root.find("svg:title", ns) or root.find("title")
    title = (title_elem.text if title_elem is not None else None) or resolved.name

    parts = [f"# {title} (SVG diagram)\n"]
    if texts:
        parts.append("## Text elements\n")
        for t in texts:
            parts.append(f"- {t}")
    else:
        parts.append("(no text elements found — this SVG may be purely graphical)")

    return "\n".join(parts)


@server.tool()
def read_image(path: str) -> str:
    """Extract text from an image using OCR.

    Requires tesseract to be installed on the system.
    For diagram understanding, consider using a multimodal model instead.

    Args:
        path: File path (relative to workspace root or absolute).
    """
    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        from PIL import Image  # type: ignore[import-not-found]  # noqa: PLC0415
    except ImportError:
        return (
            "ERROR: Pillow is not installed. Install it with:\n"
            "  pip install Pillow\n"
            "  # or: uv pip install Pillow"
        )

    try:
        import pytesseract  # type: ignore[import-not-found]  # noqa: PLC0415
    except ImportError:
        return (
            "ERROR: pytesseract is not installed. Install it with:\n"
            "  pip install pytesseract\n"
            "  # Also install tesseract OCR: sudo apt install tesseract-ocr"
        )

    try:
        img = Image.open(resolved)
        text = pytesseract.image_to_string(img)
    except Exception as exc:
        return f"ERROR: cannot process image: {exc}"

    if not text.strip():
        return (
            f"# {resolved.name}\n\n"
            "(no text detected via OCR — this image may contain diagrams "
            "or graphics that require a multimodal model to understand)"
        )

    return f"# {resolved.name} (OCR extraction)\n\n{text.strip()}"


@server.tool()
def read_text(path: str, start_line: int = 1, end_line: int | None = None) -> str:
    """Read a plain text file with line numbers.

    Works with any text-based format: .txt, .md, .yaml, .json, .xml, .html, etc.

    Args:
        path: File path (relative to workspace root or absolute).
        start_line: First line to include (1-indexed, default 1).
        end_line: Last line to include (inclusive). Defaults to start_line + 999.
    """
    resolved = _resolve_path(path)
    if not resolved.exists():
        return f"ERROR: file not found: {resolved}"

    try:
        content = resolved.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        return f"ERROR: cannot read file: {exc}"

    lines = content.splitlines()
    total = len(lines)
    start = max(0, start_line - 1)
    end = min(total, end_line or start + 1000)

    numbered = [f"{i + 1:>6}\t{lines[i]}" for i in range(start, end)]
    result = f"# {resolved.name} ({total} lines)\n\n" + "\n".join(numbered)

    if end < total and end_line is None:
        result += _truncation_notice(total, end - start, "lines", "start_line/end_line")
    return result


@server.tool()
def list_files(
    directory: str = ".",
    pattern: str = "*",
    recursive: bool = False,
) -> str:
    """List files in a directory, optionally filtered by glob pattern.

    Useful for discovering documents before reading them.

    Args:
        directory: Directory path (relative to workspace root or absolute).
        pattern: Glob pattern to filter files (default "*").
        recursive: Search subdirectories (default false).
    """
    resolved = _resolve_path(directory)
    if not resolved.exists():
        return f"ERROR: directory not found: {resolved}"
    if not resolved.is_dir():
        return f"ERROR: not a directory: {resolved}"

    files = sorted(resolved.rglob(pattern) if recursive else resolved.glob(pattern))

    if not files:
        return f"No files matching '{pattern}' in {resolved}"

    max_files = 200
    parts = [f"# Files in {resolved} (pattern: {pattern})\n"]
    for f in files[:max_files]:
        if f.is_file():
            size = f.stat().st_size
            rel = f.relative_to(resolved)
            if size > 1024 * 1024:
                size_str = f"{size / (1024 * 1024):.1f}MB"
            elif size > 1024:
                size_str = f"{size / 1024:.1f}KB"
            else:
                size_str = f"{size}B"
            parts.append(f"- {rel} ({size_str})")

    if len(files) > max_files:
        parts.append(f"\n... and {len(files) - max_files} more files")
    return "\n".join(parts)


def run_server(workspace: Path | None = None) -> None:
    """Entry point for the CLI launcher and ``__main__``."""
    global _workspace_root  # noqa: PLW0603
    _workspace_root = workspace or Path(os.environ.get("SWARMKIT_WORKSPACE", ".")).resolve()
    server.run()
